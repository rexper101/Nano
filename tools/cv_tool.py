"""
CV Tool
========
Reads your existing CV, updates sections using the LLM,
and saves the updated version as a .docx and .txt file.

Examples:
  "update my CV with a new project: built a Nano AI assistant"
  "add Python to my skills in CV"
  "export my CV as PDF"
  "show my CV"
"""

import os
import re
import httpx
from pathlib import Path


OLLAMA_URL = "http://localhost:11434/api/chat"
MODEL      = "qwen2.5:7b"
CV_DIR     = Path(os.path.expanduser("~/Documents/Nano_CV"))

CV_UPDATE_SYSTEM = """You are a professional CV writer.
The user will give you their current CV text and an update instruction.
Rewrite the CV incorporating the update professionally.
Keep the same sections. Output only the updated CV text, no commentary."""


class CVTool:
    def run(self, user_text: str) -> str:
        text = user_text.lower()

        if any(w in text for w in ["show", "read", "display", "view"]):
            return self._show_cv()

        if any(w in text for w in ["export", "pdf", "save"]):
            return self._export_cv()

        # Default: update the CV
        return self._update_cv(user_text)

    def _get_cv_path(self) -> Path:
        CV_DIR.mkdir(parents=True, exist_ok=True)
        # Look for existing CV
        for name in ["cv.txt", "resume.txt", "CV.txt", "Resume.txt"]:
            p = CV_DIR / name
            if p.exists():
                return p
        # Also check Desktop and Documents
        for folder in [Path(os.path.expanduser("~/Desktop")),
                       Path(os.path.expanduser("~/Documents"))]:
            for name in ["cv.txt", "resume.txt", "CV.docx", "Resume.docx"]:
                p = folder / name
                if p.exists():
                    return p
        return CV_DIR / "cv.txt"

    def _load_cv(self) -> str:
        path = self._get_cv_path()
        if path.exists():
            return path.read_text(encoding="utf-8", errors="ignore")
        # Return a blank CV template
        return self._blank_template()

    def _blank_template(self) -> str:
        return """NAME: [Your Name]
EMAIL: [your.email@example.com]
PHONE: [+91 XXXXXXXXXX]
LINKEDIN: [linkedin.com/in/yourprofile]

OBJECTIVE:
[Your career objective here]

EDUCATION:
MCA (Data Science) — [Your University] — 2023-2025

SKILLS:
Python, Machine Learning, Data Science, SQL, Git

PROJECTS:
- Nano AI Desktop Assistant — Built an offline AI assistant with voice, vision, and automation

EXPERIENCE:
[Your work experience here]

CERTIFICATIONS:
[Your certifications here]"""

    def _show_cv(self) -> str:
        cv = self._load_cv()
        print(f"\033[36m\n--- YOUR CV ---\n{cv}\n---\033[0m")
        return "Displaying your CV in the terminal."

    def _update_cv(self, instruction: str) -> str:
        current_cv = self._load_cv()
        print(f"\033[35m[CV] Updating with: {instruction[:60]}\033[0m")

        try:
            resp = httpx.post(
                OLLAMA_URL,
                json={
                    "model": MODEL,
                    "messages": [
                        {"role": "system", "content": CV_UPDATE_SYSTEM},
                        {"role": "user",   "content":
                            f"Current CV:\n{current_cv}\n\nUpdate instruction: {instruction}"},
                    ],
                    "stream": False,
                    "options": {"temperature": 0.3, "num_predict": 1024},
                },
                timeout=60.0,
            )
            resp.raise_for_status()
            updated = resp.json()["message"]["content"].strip()

            # Save updated CV
            path = CV_DIR / "cv.txt"
            path.write_text(updated, encoding="utf-8")

            # Try to also save as .docx
            self._save_docx(updated)

            return f"CV updated and saved to {path}"

        except Exception as e:
            return f"CV update failed: {e}"

    def _export_cv(self) -> str:
        cv_text = self._load_cv()
        self._save_docx(cv_text)
        path = CV_DIR / "cv_export.docx"
        # Open the file
        try:
            os.startfile(str(path))
        except Exception:
            pass
        return f"CV exported to {path}"

    def _save_docx(self, text: str):
        try:
            from docx import Document
            doc  = Document()
            doc.add_heading("Curriculum Vitae", 0)
            for line in text.split("\n"):
                if line.strip():
                    if line.isupper() or line.endswith(":"):
                        doc.add_heading(line, level=2)
                    else:
                        doc.add_paragraph(line)
            path = CV_DIR / "cv_export.docx"
            doc.save(str(path))
        except ImportError:
            pass  # python-docx not installed, skip
